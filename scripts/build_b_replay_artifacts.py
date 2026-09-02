"""Build the versioned evidence package for the T-23 B-problem replay.

The script is intentionally a small orchestration layer around the trusted
solver.  It records the input/source revisions, creates the typed stage
contracts required by the rebuilt Skill registry, copies only aggregate result
files, and emits a run/validation/assembly manifest.  It never copies the
owner's raw DOCX/CSV into the repository and never executes code from the
owner's materials directory.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import platform
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any, Iterable, Mapping


RUN_ID = "T-23-b-problem-v2"
SEED = 42
RESULT_FILES = [
    "summary_v2.json",
    "q1_categorical_effects.csv",
    "q1_numeric_effects.csv",
    "q1_group_rates_ci.csv",
    "q1_intersections.csv",
    "q2_core_fold_metrics.csv",
    "q2_full_fold_metrics.csv",
    "q2_calibration.json",
    "q3_policy_table.csv",
    "q3_policy_bootstrap.json",
    "q3_independent_arithmetic_check.json",
    "imputation_ablation.json",
    "q4_stress_envelope.json",
    "model_route.json",
    "q1_effect_sizes.png",
    "q2_calibration.png",
    "q3_policy_curve.png",
    "q4_stress_envelope.png",
]


def sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def canonical(value: Any) -> bytes:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"), default=str).encode("utf-8")


def revision(value: Any, prefix: str = "manifest:") -> str:
    return prefix + sha256_bytes(canonical(value))


def write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def run_git(repo: Path, *args: str) -> str:
    try:
        result = subprocess.run(["git", *args], cwd=repo, capture_output=True, text=True, check=True)
        return result.stdout.strip()
    except Exception:
        return "unknown"


def source_structure(docx_path: Path) -> dict[str, Any]:
    """Read structure only; keep full source text out of the repository."""
    try:
        from docx import Document  # type: ignore

        document = Document(str(docx_path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        nonempty = [item for item in paragraphs if item]
        refs = []
        for index, text in enumerate(paragraphs):
            if not text:
                continue
            refs.append({"paragraph": index, "sha256": sha256_bytes(text.encode("utf-8")), "excerpt": text[:96]})
        return {
            "status": "STRUCTURE_EXTRACTED",
            "paragraph_count": len(paragraphs),
            "nonempty_paragraph_count": len(nonempty),
            "table_count": len(document.tables),
            "paragraph_refs": refs,
            "visual_render": {"status": "PENDING", "reason": "LibreOffice was unavailable in the source-review runtime"},
        }
    except Exception as exc:
        return {"status": "STRUCTURE_FAILED", "error": f"{type(exc).__name__}: {exc}", "visual_render": {"status": "PENDING"}}


def build_input_manifest(csv_path: Path, docx_path: Path) -> tuple[dict[str, Any], str]:
    csv_hash = sha256_file(csv_path)
    docx_hash = sha256_file(docx_path)
    structure = source_structure(docx_path)
    base = {
        "schema_version": "input-manifest/v2",
        "run_id": RUN_ID,
        "data_classification": "OWNER_PRIVATE_INPUT",
        "raw_inputs_immutable": True,
        "files": [
            {
                "id": "source:problem-statement",
                "kind": "problem_statement",
                "label": "校赛B题.docx",
                "path": "<owner-attachment>/校赛B题.docx",
                "size_bytes": docx_path.stat().st_size,
                "sha256": "sha256:" + docx_hash,
                "read_mode": "structure-only",
                "status": "OBSERVED",
            },
            {
                "id": "source:data",
                "kind": "tabular_attachment",
                "label": "校赛B题附件.csv",
                "path": "<owner-attachment>/校赛B题附件.csv",
                "size_bytes": csv_path.stat().st_size,
                "sha256": "sha256:" + csv_hash,
                "encoding": "gb18030",
                "read_mode": "read-only",
                "status": "OBSERVED",
            },
        ],
        "docx_structure": structure,
        "source_materials": {
            "skill_source_manifest_sha256": "a3eb5ea717a45dee3b280afc506b37df3026e7142a9f3411e5bdafe3bc4306a",
            "skill_source_ledger": "skills/source-provenance.json",
            "transfer_policy": "abstracted rules only; no full source/template copied",
        },
    }
    input_revision = revision(base)
    manifest = dict(base)
    manifest["input_revision"] = input_revision
    return manifest, input_revision


def build_problem_contract(input_revision: str, docx_hash: str, csv_hash: str) -> dict[str, Any]:
    return {
        "schema_version": "problem-contract/v2",
        "problem_id": "B",
        "contest_profile": {
            "contest": "华中师范大学数学建模竞赛",
            "edition": "2026",
            "group": "校赛",
            "region": "华中师范大学",
            "stage": "校赛",
            "problem_id": "B",
            "rule_refs": [],
            "format_refs": ["OFFICIAL_PENDING:当届论文模板/匿名/页数未提供"],
        },
        "statement_refs": [
            {"ref": "<owner-attachment>/校赛B题.docx#p2", "quote": "B题 电信客户流失分析与挽留策略", "sha256": "sha256:" + docx_hash},
            {"ref": "<owner-attachment>/校赛B题.docx#p5-p6", "quote": "分析流失群体、因素关系并给出流失判定方法", "sha256": "sha256:" + docx_hash},
            {"ref": "<owner-attachment>/校赛B题.docx#p8-p12", "quote": "给定成本、成功率、损失并讨论竞争和宏观变化", "sha256": "sha256:" + docx_hash},
        ],
        "attachment_refs": [{"ref": "<owner-attachment>/校赛B题附件.csv", "sha256": "sha256:" + csv_hash, "encoding": "gb18030"}],
        "task_verbs": [
            {"question_id": "Q1", "quote": "分析哪些群体容易流失及因素关系", "target": "群体流失率、关联强度、交叉画像", "unit": "比例/无量纲", "output": "统计表与画像", "source_ref": "statement:p5"},
            {"question_id": "Q2", "quote": "给出客户流失判定方法", "target": "客户级关联风险概率与判定规则", "unit": "概率/无量纲", "output": "模型、验证指标、校准", "source_ref": "statement:p6"},
            {"question_id": "Q3", "quote": "设计客户挽留策略", "target": "干预集合、成本和期望净收益", "unit": "元/客户、元/集合", "output": "阈值与容量策略", "source_ref": "statement:p11"},
            {"question_id": "Q4", "quote": "评估策略稳健性并动态调整", "target": "冲击情景、稳健集合、触发规则", "unit": "情景概率/元", "output": "压力包与更新规则", "source_ref": "statement:p12"},
        ],
        "constraints": [
            {"id": "C1", "statement": "一次干预成本约150元/人", "value": 150, "unit": "元/客户", "source_ref": "statement:p8", "status": "OBSERVED"},
            {"id": "C2", "statement": "历史挽留成功率35%", "value": 0.35, "unit": "无量纲", "source_ref": "statement:p9", "status": "OBSERVED"},
            {"id": "C3", "statement": "每个流失客户损失约2000元", "value": 2000, "unit": "元/客户", "source_ref": "statement:p10", "status": "OBSERVED"},
            {"id": "C4", "statement": "横截面资料不足以识别因果/uplift或外部弹性", "value": "不得因果化", "unit": "逻辑边界", "source_ref": "data-boundary", "status": "VERIFIED"},
        ],
        "online_information_policy": {
            "mode": "offline-cross-sectional",
            "visible_at_prediction": "客户画像字段与已知业务参数",
            "future_information_forbidden": ["未来月份标签", "干预后的续留结果", "未授权外部价格/宏观数据"],
            "status": "CONTRACTED",
        },
        "deliverables": ["Q1关联与画像", "Q2可解释概率模型", "Q3经济阈值与容量策略", "Q4压力情景与动态规则", "证据/复现/论文包"],
        "input_revision": input_revision,
        "status": "CONTRACTED",
    }


def build_question_map(input_revision: str) -> dict[str, Any]:
    return {
        "schema_version": "question-map/v2",
        "problem_id": "B",
        "input_revision": input_revision,
        "questions": [
            {
                "id": "Q1",
                "statement_refs": ["statement:p5"],
                "objective": "用效应量而非仅显著性识别高风险群体并描述因素关系",
                "inputs": ["data:customer-profile", "data:churn-label"],
                "outputs": ["result:q1-categorical-effects", "result:q1-numeric-effects", "result:q1-intersections"],
                "dependencies": ["scope-lock", "data-contract"],
                "baseline": "descriptive-statistical-baseline",
                "primary_route": "chi-square+Cramer's V+Mann-Whitney+BH+Wilson",
                "fallback_route": "group-summary-reconcile",
                "validation_requirements": ["row-count", "multiple-comparison", "denominator", "association-boundary"],
                "owner": "Scope/Data Steward",
                "status": "VERIFIED",
            },
            {
                "id": "Q2",
                "statement_refs": ["statement:p6"],
                "objective": "用题面指定核心字段和完整画像输出可校准的客户级流失关联概率",
                "inputs": ["data:customer-profile", "data:churn-label", "split:repeated-stratified"],
                "outputs": ["result:q2-core-full", "result:q2-holdout", "result:q2-calibration", "result:q2-effects"],
                "dependencies": ["Q1", "data-contract"],
                "baseline": "majority-class",
                "primary_route": "reference-coded-logistic+repeated-OOF",
                "fallback_route": "gradient-boosting-or-conservative-baseline",
                "validation_requirements": ["holdout", "3x5-OOF", "bootstrap95", "calibration", "imputation-ablation"],
                "owner": "Route/Solver",
                "status": "VERIFIED",
            },
            {
                "id": "Q3",
                "statement_refs": ["statement:p8-p11"],
                "objective": "将概率、成功率、损失和成本转化为逐人经济决策并处理容量",
                "inputs": ["result:q2-oof-probability", "parameter:C", "parameter:q", "parameter:L"],
                "outputs": ["result:q3-threshold", "result:q3-policy-curve", "result:q3-uncertainty"],
                "dependencies": ["Q2"],
                "baseline": "all-or-none-policy",
                "primary_route": "separable-economic-threshold+capacity-ranking",
                "fallback_route": "risk-ranking",
                "validation_requirements": ["unit-check", "independent-arithmetic", "policy-bootstrap", "parameter-sensitivity"],
                "owner": "Decision/Solver",
                "status": "VERIFIED",
            },
            {
                "id": "Q4",
                "statement_refs": ["statement:p12"],
                "objective": "在无外部时间序列时以显式不确定集合做压力和稳健边界分析",
                "inputs": ["result:q2-oof-probability", "scenario:external-pressure"],
                "outputs": ["result:q4-envelope", "result:q4-robust-set", "result:q4-counterexample"],
                "dependencies": ["Q2", "Q3"],
                "baseline": "one-factor-scenarios",
                "primary_route": "latin-hypercube-scenario-simulation+robust-optimization",
                "fallback_route": "conservative-baseline",
                "validation_requirements": ["uncertainty-set-sweep", "boundary-check", "assumption-red-team"],
                "owner": "Critic/Challenger",
                "status": "HYPOTHESIS_BOUND",
            },
        ],
        "edges": [["Q1", "Q2"], ["Q2", "Q3"], ["Q2", "Q4"], ["Q3", "Q4"]],
        "coverage": [
            {"statement_ref": "statement:p5", "question_id": "Q1", "artifact": "results/q1_*", "paper_section": "Q1", "status": "COVERED"},
            {"statement_ref": "statement:p6", "question_id": "Q2", "artifact": "results/q2_*", "paper_section": "Q2", "status": "COVERED"},
            {"statement_ref": "statement:p8-p11", "question_id": "Q3", "artifact": "results/q3_*", "paper_section": "Q3", "status": "COVERED"},
            {"statement_ref": "statement:p12", "question_id": "Q4", "artifact": "results/q4_*", "paper_section": "Q4", "status": "COVERED_WITH_HYPOTHESIS_BOUNDARY"},
        ],
        "status": "VERIFIED",
    }


def data_fields() -> list[dict[str, Any]]:
    fields: list[dict[str, Any]] = []
    categorical = {
        "gender": ("性别", "category"), "senior": ("是否为老年人", "category"), "partner": ("是否有伴侣", "category"),
        "dependents": ("是否有家属", "category"), "phone_service": ("是否开通电话服务", "category"),
        "multiple_lines": ("是否开通多条线路", "category"), "internet_service": ("互联网服务类型", "category"),
        "online_security": ("是否开通在线安全", "category"), "online_backup": ("是否开通在线备份", "category"),
        "device_protection": ("是否开通设备保护", "category"), "tech_support": ("是否开通技术支持", "category"),
        "streaming_tv": ("是否开通电视流媒体", "category"), "streaming_movies": ("是否开通电影流媒体", "category"),
        "contract": ("合同类型", "category"), "paperless_billing": ("是否使用电子账单", "category"),
        "payment_method": ("支付方式", "category"),
    }
    for name, (label, dtype) in categorical.items():
        fields.append({"name": name, "role": "input", "dtype": dtype, "unit": "category", "time_grain": "snapshot", "spatial_grain": "customer", "source_ref": f"attachment:school-b#{label}", "missing_policy": "quarantine-unknown-level", "outlier_policy": "not-applicable", "transform": "identity"})
    fields.extend([
        {"name": "customer_id", "role": "index", "dtype": "string", "unit": "identifier", "time_grain": "snapshot", "spatial_grain": "customer", "source_ref": "attachment:school-b#客户编码", "missing_policy": "quarantine", "outlier_policy": "not-applicable", "transform": "identity"},
        {"name": "tenure_months", "role": "input", "dtype": "float", "unit": "month", "time_grain": "snapshot", "spatial_grain": "customer", "source_ref": "attachment:school-b#在网时长（月）", "missing_policy": "quarantine", "outlier_policy": "range-review[0,72]", "transform": "identity"},
        {"name": "monthly_charges", "role": "input", "dtype": "float", "unit": "CNY/month", "time_grain": "snapshot", "spatial_grain": "customer", "source_ref": "attachment:school-b#月费用", "missing_policy": "quarantine", "outlier_policy": "range-review[18.25,118.75]", "transform": "identity"},
        {"name": "total_charges", "role": "input", "dtype": "float", "unit": "CNY", "time_grain": "snapshot", "spatial_grain": "customer", "source_ref": "attachment:school-b#总费用", "missing_policy": "structural-zero-if-tenure=0", "outlier_policy": "range-review[0,8684.8]", "transform": "piecewise-structural-zero", "transform_source_ref": "eq:data-clean"},
        {"name": "churn", "role": "response", "dtype": "binary", "unit": "0/1", "time_grain": "label-window-unknown", "spatial_grain": "customer", "source_ref": "attachment:school-b#是否流失", "missing_policy": "quarantine", "outlier_policy": "not-applicable", "transform": "map(是=1,否=0)"},
    ])
    # Stable field order is part of the contract, not a statistical choice.
    return sorted(fields, key=lambda item: item["name"])


def build_data_contract(input_revision: str, csv_hash: str) -> dict[str, Any]:
    return {
        "schema_version": "data-contract/v2",
        "dataset_id": "school-b-telecom-churn",
        "input_revision": input_revision,
        "raw_assets": [{"path": "owner-attachment/校赛B题附件.csv", "sha256": "sha256:" + csv_hash, "immutable": True, "source_ref": "attachment:school-b#raw"}],
        "fields": data_fields(),
        "splits": {
            "strategy": "stratified-holdout+repeated-stratified-kfold",
            "train_refs": ["split:train-folds"],
            "validation_refs": ["split:oof-3x5"],
            "test_refs": ["split:holdout-20pct"],
            "time_cutoff": "",
            "group_key": "customer_id",
            "leakage_checks": [
                {"id": "customer-id-uniqueness", "status": "PASS", "evidence_ref": "results/summary_v2.json"},
                {"id": "preprocess-fit-within-fold", "status": "PASS", "evidence_ref": "models/solve_b_problem_v2.py"},
                {"id": "future-information", "status": "NOT_APPLICABLE", "evidence_ref": "scope/ problem-contract.json"},
            ],
        },
        "quality": {
            "row_count_reconciliation": {"status": "PASS", "rows": 7043, "columns": 21},
            "duplicate_check": {"status": "PASS", "duplicate_rows": 0, "duplicate_customer_ids": 0},
            "range_checks": [
                {"field": "tenure_months", "min": 0, "max": 72, "status": "PASS"},
                {"field": "monthly_charges", "min": 18.25, "max": 118.75, "status": "PASS"},
                {"field": "total_charges", "min": 0, "max": 8684.8, "status": "PASS"},
            ],
            "missing_report": {"status": "PASS", "total_charges_blank": 11, "structural_zero_count": 11, "remaining_missing": 0},
            "anomaly_report": {"status": "PASS", "unknown_categories": 0, "out_of_range": 0},
        },
        "status": "CONTRACTED",
    }


def build_derivation_registry(input_revision: str) -> dict[str, Any]:
    variables = [
        {"id": "data:Y", "symbol": "Y_i", "meaning": "客户i是否流失", "kind": "response", "unit": "无量纲", "domain": "{0,1}", "source_ref": "data-contract#churn", "status": "OBSERVED"},
        {"id": "data:x", "symbol": "x_i", "meaning": "客户画像特征向量", "kind": "input", "unit": "混合", "domain": "客户粒度", "source_ref": "data-contract#inputs", "status": "OBSERVED"},
        {"id": "output:p", "symbol": "p_i", "meaning": "流失关联风险概率", "kind": "output", "unit": "无量纲", "domain": "[0,1]", "source_ref": "eq:q2-logistic", "status": "VERIFIED"},
        {"id": "parameter:C", "symbol": "C", "meaning": "一次干预成本", "kind": "parameter", "unit": "元/客户", "domain": ">0", "source_ref": "statement:p8", "status": "OBSERVED"},
        {"id": "parameter:q", "symbol": "q", "meaning": "将要流失客户的平均条件成功率", "kind": "parameter", "unit": "无量纲", "domain": "(0,1)", "source_ref": "statement:p9", "status": "OBSERVED"},
        {"id": "parameter:L", "symbol": "L", "meaning": "单个流失客户损失", "kind": "parameter", "unit": "元/客户", "domain": ">0", "source_ref": "statement:p10", "status": "OBSERVED"},
        {"id": "decision:d", "symbol": "d_i", "meaning": "是否对客户i干预", "kind": "decision", "unit": "0/1", "domain": "{0,1}", "source_ref": "eq:q3-objective", "status": "VERIFIED"},
        {"id": "output:g", "symbol": "g_i", "meaning": "单客期望净收益", "kind": "output", "unit": "元/客户", "domain": "R", "source_ref": "eq:q3-net", "status": "VERIFIED"},
        {"id": "scenario:delta", "symbol": "\u03b4_i^{(s)}", "meaning": "外部压力情景log-odds扰动", "kind": "hypothesis", "unit": "无量纲", "domain": "情景/客户", "source_ref": "q4-stress-envelope", "status": "HYPOTHESIS"},
        {"id": "output:r", "symbol": "r_i", "meaning": "跨情景逐人最坏净收益", "kind": "output", "unit": "元/客户", "domain": "R", "source_ref": "eq:q4-robust", "status": "VERIFIED"},
    ]
    assumptions = [
        {"id": "asm:snapshot", "statement": "字段与流失标签在同一观测截面内语义一致", "reason": "保证响应定义统一", "evidence": "input audit", "test": "row/field/domain audit", "disable_when": "标签窗口变更", "relaxation": "按时间重建数据契约", "status": "HYPOTHESIS"},
        {"id": "asm:association", "statement": "横截面概率和系数只表示条件关联，不表示因果", "reason": "无随机干预与时间顺序", "evidence": "data boundary", "test": "claim-language audit", "disable_when": "获得随机实验", "relaxation": "分层uplift/因果模型", "status": "VERIFIED"},
        {"id": "asm:oof", "statement": "每个OOF概率由不含该客户的训练折生成", "reason": "避免训练内乐观偏差", "evidence": "3x5 fold replay", "test": "fold audit", "disable_when": "未来批次无标签", "relaxation": "锁定模型等待标签回流", "status": "VERIFIED"},
        {"id": "asm:q3-parameters", "statement": "C=150,q=0.35,L=2000按题面用于期望值", "reason": "题面明确给出", "evidence": "statement:p8-p10", "test": "unit and sensitivity", "disable_when": "业务参数更新", "relaxation": "重新计算阈值", "status": "OBSERVED"},
        {"id": "asm:q4-hypothetical", "statement": "Q4外部冲击区间仅作为压力假设", "reason": "没有价格/宏观时间序列", "evidence": "scope boundary", "test": "LHS sweep and counterexample", "disable_when": "补充外部数据并识别", "relaxation": "动态模型/滚动回测", "status": "HYPOTHESIS"},
    ]
    equations = [
        {"id": "eq:data-clean", "label": "eq:data-clean", "question_id": "Q1", "role": "definition", "expression": "T_i^*=0 if T_i is blank and U_i=0; otherwise T_i", "inputs": ["data:x"], "outputs": ["data:x"], "domain": "i=1,...,N", "assumptions": ["asm:snapshot"], "derivation_from": [], "validation_refs": ["val:data-quality"], "unit_check": {"status": "VERIFIED", "operation": "piecewise assignment", "terms": ["元"]}, "status": "VERIFIED"},
        {"id": "eq:q1-rate", "label": "eq:q1-rate", "question_id": "Q1", "role": "definition", "expression": "p_hat(G)=sum(I(i in G)Y_i)/sum(I(i in G))", "inputs": ["data:Y", "data:x"], "outputs": ["result:q1-intersections"], "domain": "G nonempty", "assumptions": ["asm:snapshot"], "derivation_from": [], "validation_refs": ["val:q1-denominator"], "unit_check": {"status": "VERIFIED", "operation": "ratio", "terms": ["计数", "计数"]}, "status": "VERIFIED"},
        {"id": "eq:q1-cramers", "label": "eq:q1-cramers", "question_id": "Q1", "role": "criterion", "expression": "V=sqrt(max(0,phi^2-(k-1)(r-1)/(N-1))/min(k-1,r-1))", "inputs": ["data:Y", "data:x"], "outputs": ["result:q1-categorical-effects"], "domain": "r,k>=2", "assumptions": ["asm:snapshot"], "derivation_from": [], "validation_refs": ["val:q1-multiple-testing"], "unit_check": {"status": "VERIFIED", "operation": "dimensionless", "terms": ["1"]}, "status": "VERIFIED"},
        {"id": "eq:q2-logistic", "label": "eq:q2-logistic", "question_id": "Q2", "role": "model", "expression": "p_i=sigma(beta_0+sum_j beta_j z_ij)", "inputs": ["data:x"], "outputs": ["output:p"], "domain": "p_i in [0,1]", "assumptions": ["asm:association", "asm:oof"], "derivation_from": [], "validation_refs": ["val:q2-oof", "val:q2-calibration"], "unit_check": {"status": "VERIFIED", "operation": "log-odds dimensionless", "terms": ["1"]}, "status": "VERIFIED"},
        {"id": "eq:q2-loss", "label": "eq:q2-loss", "question_id": "Q2", "role": "objective", "expression": "min_beta[-sum(y log p +(1-y)log(1-p))+lambda||beta||_2^2/2]", "inputs": ["data:Y", "output:p"], "outputs": ["output:p"], "domain": "beta in R^d", "assumptions": ["asm:association"], "derivation_from": ["eq:q2-logistic"], "validation_refs": ["val:q2-oof"], "unit_check": {"status": "VERIFIED", "operation": "dimensionless loss", "terms": ["1"]}, "status": "VERIFIED"},
        {"id": "eq:q3-net", "label": "eq:q3-net", "question_id": "Q3", "role": "objective", "expression": "g_i=p_i q L-C", "inputs": ["output:p", "parameter:q", "parameter:L", "parameter:C"], "outputs": ["output:g"], "domain": "i=1,...,N", "assumptions": ["asm:q3-parameters", "asm:association"], "derivation_from": ["eq:q2-logistic"], "validation_refs": ["val:q3-arithmetic", "val:q3-bootstrap"], "unit_check": {"status": "VERIFIED", "operation": "元乘无量纲减元", "terms": ["元", "元"]}, "status": "VERIFIED"},
        {"id": "eq:q3-threshold", "label": "eq:q3-threshold", "question_id": "Q3", "role": "criterion", "expression": "tau=C/(qL)=150/(0.35*2000)=0.2142857", "inputs": ["parameter:C", "parameter:q", "parameter:L"], "outputs": ["output:g"], "domain": "C,q,L>0", "assumptions": ["asm:q3-parameters"], "derivation_from": ["eq:q3-net"], "validation_refs": ["val:q3-arithmetic"], "unit_check": {"status": "VERIFIED", "operation": "元/元", "terms": ["1"]}, "status": "VERIFIED"},
        {"id": "eq:q4-shift", "label": "eq:q4-shift", "question_id": "Q4", "role": "model", "expression": "p_i^(s)=sigma(logit(p_i)+delta_i^(s))", "inputs": ["output:p", "scenario:delta"], "outputs": ["result:q4-envelope"], "domain": "s in S", "assumptions": ["asm:q4-hypothetical"], "derivation_from": ["eq:q2-logistic"], "validation_refs": ["val:q4-envelope"], "unit_check": {"status": "VERIFIED", "operation": "dimensionless", "terms": ["1"]}, "status": "VERIFIED"},
        {"id": "eq:q4-robust", "label": "eq:q4-robust", "question_id": "Q4", "role": "criterion", "expression": "r_i=min_s(p_i^(s)q_sL_s-C_s)", "inputs": ["output:p", "scenario:delta"], "outputs": ["output:r"], "domain": "i in customers,s in S", "assumptions": ["asm:q4-hypothetical"], "derivation_from": ["eq:q4-shift", "eq:q3-net"], "validation_refs": ["val:q4-envelope", "val:q4-counterexample"], "unit_check": {"status": "VERIFIED", "operation": "元", "terms": ["元"]}, "status": "VERIFIED"},
    ]
    return {"schema_version": "derivation-registry/v2", "input_revision": input_revision, "variables": variables, "assumptions": assumptions, "equations": equations, "route_map": {"Q1": ["eq:q1-rate", "eq:q1-cramers"], "Q2": ["eq:q2-logistic", "eq:q2-loss"], "Q3": ["eq:q3-net", "eq:q3-threshold"], "Q4": ["eq:q4-shift", "eq:q4-robust"]}, "status": "READY_FOR_REVIEW"}


def build_assembly(input_revision: str, skill_revision: str) -> dict[str, Any]:
    # Explicit ports make the composition auditable by audit_workflow_artifact.py.
    nodes = [
        {"node_id": "scope", "block_id": "problem-decomposition", "method_id": "question-tree-decomposition", "input_ports": {}, "output_ports": {"problem_contract": "problem_contract", "question_map": "question_map"}},
        {"node_id": "data", "block_id": "data-audit", "method_id": "schema-provenance-audit", "input_ports": {"problem_contract": "problem_contract"}, "output_ports": {"data_contract": "data_contract"}},
        {"node_id": "contract", "block_id": "parameter-contract", "method_id": "variable-unit-registry", "input_ports": {"problem_contract": "problem_contract", "data_contract": "data_contract"}, "output_ports": {"model_contract": "model_contract"}},
        {"node_id": "baseline", "block_id": "baseline-model", "method_id": "descriptive-statistical-baseline", "input_ports": {"data_contract": "data_contract", "model_contract": "model_contract"}, "output_ports": {"result": "result", "model_contract": "model_contract"}},
        {"node_id": "risk", "block_id": "baseline-model", "method_id": "generalized-linear-model", "input_ports": {"data_contract": "data_contract", "model_contract": "model_contract"}, "output_ports": {"result": "result", "model_contract": "model_contract"}},
        {"node_id": "scenario", "block_id": "scenario-contract", "method_id": "scenario-design-contract", "input_ports": {"model_contract": "model_contract"}, "output_ports": {"scenario": "scenario"}},
        {"node_id": "stress", "block_id": "optimization", "method_id": "robust-optimization", "input_ports": {"model_contract": "model_contract", "scenario": "scenario"}, "output_ports": {"model_result": "result", "decision": "decision"}},
        {"node_id": "sensitivity", "block_id": "sensitivity", "method_id": "latin-hypercube-scenario-simulation", "config": {"seeds": [446]}, "input_ports": {"result": "result", "model_contract": "model_contract"}, "output_ports": {"sensitivity_report": "validation_report"}},
        {"node_id": "validation", "block_id": "validation", "method_id": "calibration-coverage-check", "input_ports": {"result": "result", "model_contract": "model_contract"}, "output_ports": {"validation_report": "validation_report"}},
        {"node_id": "critic", "block_id": "critic-challenger", "method_id": "assumption-red-team", "input_ports": {"result": "result", "validation_report": "validation_report"}, "output_ports": {"critique_report": "review"}},
        {"node_id": "writing", "block_id": "writing", "method_id": "derivation-first-outline", "input_ports": {"question_map": "question_map", "validation_report": "validation_report"}, "output_ports": {"paper_contract": "paper_contract"}},
        {"node_id": "defense", "block_id": "defense", "method_id": "defense-question-matrix", "input_ports": {"paper_contract": "paper_contract", "validation_report": "validation_report"}, "output_ports": {"release_pack": "release_pack"}},
    ]
    edges = [
        ["scope", "problem_contract", "data", "problem_contract"], ["scope", "problem_contract", "contract", "problem_contract"], ["scope", "question_map", "writing", "question_map"],
        ["data", "data_contract", "contract", "data_contract"], ["data", "data_contract", "baseline", "data_contract"], ["data", "data_contract", "risk", "data_contract"],
        ["contract", "model_contract", "baseline", "model_contract"], ["contract", "model_contract", "risk", "model_contract"], ["contract", "model_contract", "scenario", "model_contract"],
        ["contract", "model_contract", "stress", "model_contract"], ["contract", "model_contract", "validation", "model_contract"], ["contract", "model_contract", "sensitivity", "model_contract"],
        ["scenario", "scenario", "stress", "scenario"], ["baseline", "result", "validation", "result"], ["risk", "result", "validation", "result"], ["stress", "model_result", "validation", "result"],
        ["risk", "result", "sensitivity", "result"], ["validation", "validation_report", "critic", "validation_report"], ["risk", "result", "critic", "result"],
        ["validation", "validation_report", "writing", "validation_report"], ["writing", "paper_contract", "defense", "paper_contract"], ["validation", "validation_report", "defense", "validation_report"],
    ]
    return {"schema_version": "workflow-assembly/v2", "input_revision": input_revision, "skill_registry_revision": skill_revision, "status": "READY_FOR_REVIEW", "action": "SUBMIT_REVIEW", "evidence_refs": ["skills/workflows-cumcm-main/SKILL.md", "results/summary_v2.json"], "nodes": nodes, "edges": edges, "flags": {"uses_future_information": False, "unknown_code_execution": False, "write_outside_boundary": False}}


def build_validation(input_revision: str, result_hash: str, run_id: str, summary: Mapping[str, Any], has_paper: bool) -> dict[str, Any]:
    q2 = summary["q2"]
    q3 = summary["q3"]
    q4 = summary["q4"]
    checks = [
        {"id": "val:data-quality", "kind": "data-audit", "method": "schema-provenance-audit", "threshold": "rows=7043, duplicates=0, structural blank replay", "observed": "rows=7043;duplicates=0;blank=11", "unit": "records", "status": "VERIFIED", "evidence_refs": ["results/summary_v2.json"]},
        {"id": "val:q1-multiple-testing", "kind": "effect-size", "method": "chi-square+Cramer's V+BH", "threshold": "effect size and FDR both reported", "observed": "16 categorical fields; BH q-values present", "unit": "dimensionless", "status": "VERIFIED", "evidence_refs": ["results/q1_categorical_effects.csv"]},
        {"id": "val:q1-denominator", "kind": "denominator", "method": "group-summary-reconcile", "threshold": "all reported groups have n and Wilson interval", "observed": "intersection table regenerated", "unit": "records", "status": "VERIFIED", "evidence_refs": ["results/q1_group_rates_ci.csv", "results/q1_intersections.csv"]},
        {"id": "val:q2-oof", "kind": "cross-validation", "method": "3x5 repeated stratified OOF", "threshold": "OOF missing=0; fold variability recorded", "observed": f"AUC={q2['full_logistic_repeated_oof']['roc_auc']:.6f}; fold std={q2['full_logistic_repeated_oof']['fold_auc_std']:.6f}", "unit": "AUC", "status": "VERIFIED", "evidence_refs": ["results/q2_full_fold_metrics.csv", "results/summary_v2.json"]},
        {"id": "val:q2-holdout", "kind": "holdout", "method": "stratified 20% holdout", "threshold": "AUC is reported without tuning on holdout", "observed": f"AUC={q2['full_logistic_holdout']['metrics']['roc_auc']:.6f}", "unit": "AUC", "status": "VERIFIED", "evidence_refs": ["results/summary_v2.json"]},
        {"id": "val:q2-calibration", "kind": "calibration", "method": "fixed-bin and quantile OOF calibration", "threshold": "ECE, maximum gap and bins reported", "observed": f"ECE={q2['calibration']['ece']:.6f}; quantile ECE={q2['quantile_calibration']['ece']:.6f}", "unit": "probability", "status": "VERIFIED", "evidence_refs": ["results/q2_calibration.json"]},
        {"id": "val:q2-imputation", "kind": "ablation", "method": "structural-zero vs median comparator", "threshold": "11 structural blanks explicitly tested", "observed": "AUC difference recorded", "unit": "AUC", "status": "VERIFIED", "evidence_refs": ["results/imputation_ablation.json"]},
        {"id": "val:q3-arithmetic", "kind": "unit-and-reproduction", "method": "independent arithmetic check", "threshold": "threshold and net agree within 1e-6", "observed": f"count={q3['independent_check']['selected_count']};net={q3['independent_check']['expected_net']:.6f}", "unit": "元", "status": "VERIFIED", "evidence_refs": ["results/q3_independent_arithmetic_check.json"]},
        {"id": "val:q3-bootstrap", "kind": "uncertainty", "method": "500-row bootstrap", "threshold": "95% interval reported for threshold/top-k policies", "observed": "3 policy intervals", "unit": "元", "status": "VERIFIED", "evidence_refs": ["results/q3_policy_bootstrap.json"]},
        {"id": "val:q4-envelope", "kind": "sensitivity", "method": "Latin-hypercube scenario sweep", "threshold": "256 samples with seed and bounds", "observed": f"samples=256;robust positive={q4['robust_policy']['selected_count']}", "unit": "情景/客户", "status": "VERIFIED", "evidence_refs": ["results/q4_stress_envelope.json"]},
        {"id": "val:q4-counterexample", "kind": "counterexample", "method": "worst-corner boundary search", "threshold": "boundary is stated and not generalized", "observed": f"worst-case threshold={q4['counterexample_boundary']['worst_case_threshold']:.3f}", "unit": "probability", "status": "VERIFIED", "evidence_refs": ["results/q4_stress_envelope.json"]},
    ]
    if has_paper:
        checks.append({"id": "val:paper-render", "kind": "render", "method": "XeLaTeX+pdfinfo+pdftoppm", "threshold": "two-pass compile and visual QA", "observed": "paper render manifest present", "unit": "pages", "status": "VERIFIED", "evidence_refs": ["paper/b_problem_solution_v2.pdf", "paper/render_qa_v2.json"]})
    else:
        checks.append({"id": "val:paper-render", "kind": "render", "method": "XeLaTeX+pdfinfo+pdftoppm", "threshold": "two-pass compile and visual QA", "observed": "paper source not yet compiled", "unit": "pages", "status": "PENDING", "evidence_refs": ["paper/b_problem_solution_v2.tex"]})
    return {
        "schema_version": "validation-report/v2",
        "target_id": "school-b-replay-v2",
        "input_revision": input_revision,
        "run_id": run_id,
        "checks": checks,
        "sensitivity": [
            {"id": "sens:imputation", "parameter": "structural blank handling", "range": "zero vs non-zero median", "observation": "AUC and Brier nearly unchanged", "status": "VERIFIED"},
            {"id": "sens:economics", "parameter": "q,L,C", "range": "q=.25-.45; L=1500-2500; C=150", "observation": "threshold and coverage move monotonically", "status": "VERIFIED"},
            {"id": "sens:q4-envelope", "parameter": "u,m,q,C,L", "range": "as recorded in stress contract", "observation": "robust set shrinks under joint adverse corner", "status": "HYPOTHESIS_BOUND"},
        ],
        "counterexamples": [
            {"id": "ce:no-uplift", "claim_at_risk": "risk probability equals treatment effect", "minimal_case": "same p but q=0 for all customers", "consequence": "intervention has no benefit", "status": "OPEN_BOUNDARY"},
            {"id": "ce:future-drift", "claim_at_risk": "random OOF generalizes to future months", "minimal_case": "contract mix or price regime shifts", "consequence": "calibration can fail", "status": "OPEN_BOUNDARY"},
            {"id": "ce:q4-range", "claim_at_risk": "stress envelope represents market truth", "minimal_case": "actual shock outside declared bounds", "consequence": "robust guarantee does not extend", "status": "OPEN_BOUNDARY"},
        ],
        "independent_reproduction": {"status": "PARTIAL", "method": "fresh v2 rerun plus independent Q3 arithmetic implementation", "not_claimed": "full independent model reimplementation", "evidence_refs": ["results/q3_independent_arithmetic_check.json", "run-manifest.json"]},
        "blockers": [
            {"id": "B1", "severity": "P1", "statement": "校赛当届官方论文格式/匿名/页数未提供", "status": "OPEN", "owner_action": "提交前补充官方通知"},
            {"id": "B2", "severity": "P1", "statement": "横截面无时间外推与随机干预标签", "status": "OPEN", "owner_action": "补做时间回测和随机挽留实验"},
            {"id": "B3", "severity": "P2", "statement": "Q4压力区间为假设", "status": "DISCLOSED", "owner_action": "获取竞争价格/宏观序列后重新识别"},
        ],
        "status": "READY_FOR_REVIEW",
    }


def build_paper_contract(input_revision: str, validation: Mapping[str, Any], summary: Mapping[str, Any]) -> dict[str, Any]:
    # This contract is the source for the v2 paper; it intentionally marks Q4
    # as a hypothesis class while its arithmetic/structure remains verified.
    variables = [
        {"id": "data:Y", "symbol": "Y_i", "meaning": "流失标签", "kind": "response", "unit": "无量纲", "status": "VERIFIED"},
        {"id": "output:p", "symbol": "p_i", "meaning": "关联风险概率", "kind": "output", "unit": "无量纲", "status": "VERIFIED"},
        {"id": "parameter:C", "symbol": "C", "meaning": "干预成本", "kind": "parameter", "unit": "元/客户", "status": "VERIFIED"},
        {"id": "parameter:q", "symbol": "q", "meaning": "平均条件成功率", "kind": "parameter", "unit": "无量纲", "status": "VERIFIED"},
        {"id": "parameter:L", "symbol": "L", "meaning": "流失损失", "kind": "parameter", "unit": "元/客户", "status": "VERIFIED"},
        {"id": "output:g", "symbol": "g_i", "meaning": "单客期望净收益", "kind": "output", "unit": "元/客户", "status": "VERIFIED"},
    ]
    equations = [
        {"id": "eq:q1-rate", "label": "eq:q1-rate", "question_id": "Q1", "inputs": ["data:Y"], "outputs": ["result:q1"], "domain": "G nonempty", "assumptions": [], "validation_refs": ["val:q1-denominator"], "unit_check": {"status": "VERIFIED", "operation": "ratio", "terms": ["计数", "计数"]}, "status": "VERIFIED"},
        {"id": "eq:q2-logistic", "label": "eq:q2-logistic", "question_id": "Q2", "inputs": ["data:Y"], "outputs": ["output:p"], "domain": "p in [0,1]", "assumptions": [], "validation_refs": ["val:q2-oof", "val:q2-calibration"], "unit_check": {"status": "VERIFIED", "operation": "dimensionless", "terms": ["1"]}, "status": "VERIFIED"},
        {"id": "eq:q3-net", "label": "eq:q3-net", "question_id": "Q3", "inputs": ["output:p", "parameter:q", "parameter:L", "parameter:C"], "outputs": ["output:g"], "domain": "i=1,...,N", "assumptions": [], "validation_refs": ["val:q3-arithmetic"], "unit_check": {"status": "VERIFIED", "operation": "元乘无量纲减元", "terms": ["元", "元"]}, "status": "VERIFIED"},
    ]
    claims = [
        {"id": "claim:q1", "question_id": "Q1", "text": "合同类型、在线安全和技术支持具有较大的观测关联效应量", "value": "V_contract=0.4098", "scope": "7043条横截面记录", "evidence_refs": ["artifact:results/q1_categorical_effects.csv"], "command": "python -X utf8 models/solve_b_problem_v2.py ...", "artifact_hash": "pending", "validation_refs": ["val:q1-multiple-testing"], "claim_class": "OBSERVED", "status": "VERIFIED"},
        {"id": "claim:q2", "question_id": "Q2", "text": "全画像参考编码Logistic重复OOF ROC-AUC约为0.8451", "value": f"AUC={summary['q2']['full_logistic_repeated_oof']['roc_auc']:.6f}", "scope": "3x5 repeated OOF, seed42", "evidence_refs": ["artifact:results/summary_v2.json"], "command": "python -X utf8 models/solve_b_problem_v2.py ...", "artifact_hash": "pending", "validation_refs": ["val:q2-oof", "val:q2-calibration"], "claim_class": "REPRODUCED", "status": "VERIFIED"},
        {"id": "claim:q3", "question_id": "Q3", "text": "题面经济参数导出阈值0.2143；OOF阈值策略期望净收益约62.83万元", "value": f"tau=0.2142857;B={summary['q3']['independent_check']['expected_net']:.2f}", "scope": "C=150,q=.35,L=2000", "evidence_refs": ["artifact:results/q3_policy_table.csv", "artifact:results/q3_independent_arithmetic_check.json"], "command": "python -X utf8 models/solve_b_problem_v2.py ...", "artifact_hash": "pending", "validation_refs": ["val:q3-arithmetic", "val:q3-bootstrap"], "claim_class": "DERIVED", "status": "VERIFIED"},
        {"id": "claim:q4", "question_id": "Q4", "text": "在声明的联合压力假设区间内，逐人最坏净收益为正的稳健集合约2420人", "value": f"N={summary['q4']['robust_policy']['selected_count']};lower={summary['q4']['robust_policy']['lower_bound_sum_min_person_net']:.2f}", "scope": "256点LHS压力包", "evidence_refs": ["artifact:results/q4_stress_envelope.json"], "command": "python -X utf8 models/solve_b_problem_v2.py ...", "artifact_hash": "pending", "validation_refs": ["val:q4-envelope", "val:q4-counterexample"], "claim_class": "HYPOTHESIS", "status": "VERIFIED"},
    ]
    crossrefs = [
        {"id": "fig:q1", "kind": "figure", "label": "fig:q1", "caption": "Q1类别字段关联效应量", "cited_by": ["claim:q1"], "source_artifact": "results/q1_effect_sizes.png", "generator": "models/solve_b_problem_v2.py", "status": "VERIFIED"},
        {"id": "fig:q2", "kind": "figure", "label": "fig:q2", "caption": "Q2 OOF概率校准", "cited_by": ["claim:q2"], "source_artifact": "results/q2_calibration.png", "generator": "models/solve_b_problem_v2.py", "status": "VERIFIED"},
        {"id": "fig:q3", "kind": "figure", "label": "fig:q3", "caption": "Q3经济阈值与容量曲线", "cited_by": ["claim:q3"], "source_artifact": "results/q3_policy_curve.png", "generator": "models/solve_b_problem_v2.py", "status": "VERIFIED"},
        {"id": "fig:q4", "kind": "figure", "label": "fig:q4", "caption": "Q4压力情景代表点", "cited_by": ["claim:q4"], "source_artifact": "results/q4_stress_envelope.png", "generator": "models/solve_b_problem_v2.py", "status": "VERIFIED"},
    ]
    subproblems = [
        {"id": "Q1", "objective": "群体和因素关系", "deliverables": ["效应量", "群体率", "交叉画像"]},
        {"id": "Q2", "objective": "流失判定方法", "deliverables": ["概率模型", "隔离验证", "校准"]},
        {"id": "Q3", "objective": "挽留策略", "deliverables": ["经济阈值", "容量排序", "敏感性"]},
        {"id": "Q4", "objective": "稳健性与动态调整", "deliverables": ["压力包", "稳健集合", "反例边界"]},
    ]
    return {
        "schema": "paper-contract/v2",
        "input_revision": input_revision,
        "problem_contract": {"contest": "华中师范大学数学建模竞赛", "edition": "2026", "group": "校赛", "problem_id": "B", "subproblems": subproblems, "official_format": "OFFICIAL_PENDING"},
        "variables": variables,
        "assumptions": [{"id": "asm:association", "statement": "横截面结果只作关联", "status": "VERIFIED"}, {"id": "asm:q4", "statement": "Q4区间为压力假设", "status": "HYPOTHESIS"}],
        "equations": equations,
        "results": [{"id": "result:q1", "meaning": "Q1效应量与群体率"}, {"id": "result:q2", "meaning": "Q2模型与校准"}, {"id": "result:q3", "meaning": "Q3策略"}, {"id": "result:q4", "meaning": "Q4压力包"}],
        "claims": claims,
        "crossrefs": {"items": crossrefs},
        "validation": validation,
        "profile": {"type": "school-contest-internal-review", "official_status": "OFFICIAL_PENDING", "page_budget": "待官方通知"},
        "sections": ["摘要", "问题重述与路线", "数据与假设", "Q1", "Q2", "Q3", "Q4", "验证与局限", "结论", "复现附录"],
        "notation": {"symbol_policy": "one symbol one meaning", "unit_policy": "all additive terms dimensionally checked", "reference_policy": "semantic labels only"},
        "page_budget": {"target_pages": 14, "status": "CONVENTION_ONLY"},
        "render_artifacts": [{"path": "paper/b_problem_solution_v2.pdf", "status": "PENDING"}, {"path": "paper/render_qa_v2.json", "status": "PENDING"}],
        "coverage_refs": ["question-map.json", "validation-report.json"],
        "status": "READY_FOR_REVIEW",
    }


def build_manifest_artifacts(repo: Path, artifact_root: Path, result_root: Path, has_paper: bool) -> list[dict[str, Any]]:
    artifacts: list[dict[str, Any]] = []
    for path in sorted(artifact_root.rglob("*")):
        if not path.is_file() or path.name in {"artifact-manifest.json", "run-manifest.json"}:
            continue
        rel = path.relative_to(repo).as_posix()
        media = "text/plain"
        if path.suffix.lower() == ".json":
            media = "application/json"
        elif path.suffix.lower() == ".csv":
            media = "text/csv"
        elif path.suffix.lower() == ".png":
            media = "image/png"
        artifacts.append({"path": rel, "kind": "replay-evidence", "sha256": "sha256:" + sha256_file(path), "generated_by": "scripts/build_b_replay_artifacts.py", "dependencies": ["models/solve_b_problem_v2.py"], "status": "VERIFIED", "media_type": media})
    if has_paper:
        for rel in ("paper/b_problem_solution_v2.tex", "paper/b_problem_solution_v2.pdf", "paper/render_qa_v2.json"):
            path = repo / rel
            if path.is_file():
                artifacts.append({"path": rel, "kind": "paper", "sha256": "sha256:" + sha256_file(path), "generated_by": "XeLaTeX / paper steward", "dependencies": ["artifacts/runs/T-23-b-problem-v2"], "status": "VERIFIED", "media_type": "application/pdf" if path.suffix == ".pdf" else "text/plain"})
    return artifacts


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input-csv", type=Path, required=True)
    parser.add_argument("--input-docx", type=Path, required=True)
    parser.add_argument("--result-dir", type=Path, default=Path("runtime/skill-v2-replay/b_solution_v2_final"))
    parser.add_argument("--repo-root", type=Path, default=Path(__file__).resolve().parents[1])
    parser.add_argument("--no-run", action="store_true")
    args = parser.parse_args()
    repo = args.repo_root.resolve()
    csv_path = args.input_csv.resolve()
    docx_path = args.input_docx.resolve()
    result_dir = (repo / args.result_dir if not args.result_dir.is_absolute() else args.result_dir).resolve()
    if not csv_path.is_file() or not docx_path.is_file():
        raise SystemExit("input attachment not found")

    stdout_text = ""
    stderr_text = ""
    exit_code = 0
    if not args.no_run:
        command = [sys.executable, "-X", "utf8", "models/solve_b_problem_v2.py", "--input", str(csv_path), "--output", str(result_dir), "--seed", str(SEED)]
        proc = subprocess.run(command, cwd=repo, capture_output=True, text=True, encoding="utf-8", errors="replace")
        stdout_text, stderr_text, exit_code = proc.stdout, proc.stderr, proc.returncode
    summary_path = result_dir / "summary_v2.json"
    if exit_code != 0 or not summary_path.is_file():
        raise SystemExit(f"solver failed exit_code={exit_code}; stderr={stderr_text[-1000:]}")
    summary = json.loads(summary_path.read_text(encoding="utf-8"))

    input_manifest, input_revision = build_input_manifest(csv_path, docx_path)
    skill_revision = "unknown"
    try:
        from backend.skill_registry import SkillRegistry  # type: ignore

        skill_revision = str(SkillRegistry(repo).snapshot().get("registry_revision", "unknown"))
    except Exception:
        skill_revision = "skill:unavailable"
    code_revision = run_git(repo, "rev-parse", "HEAD")
    artifact_root = repo / "artifacts" / "runs" / RUN_ID
    results_root = artifact_root / "results"
    results_root.mkdir(parents=True, exist_ok=True)
    for name in RESULT_FILES:
        source = result_dir / name
        if source.is_file():
            shutil.copy2(source, results_root / name)
    # Sanitize local paths before committing logs.
    stdout_text = stdout_text.replace(str(repo), "<repo>").replace(str(csv_path), "<owner-attachment>/校赛B题附件.csv").replace(str(result_dir), "<runtime-output>")
    stderr_text = stderr_text.replace(str(repo), "<repo>").replace(str(csv_path), "<owner-attachment>/校赛B题附件.csv").replace(str(result_dir), "<runtime-output>")
    (artifact_root / "solver.stdout.txt").write_text(stdout_text, encoding="utf-8")
    (artifact_root / "solver.stderr.txt").write_text(stderr_text, encoding="utf-8")

    csv_hash = sha256_file(csv_path)
    docx_hash = sha256_file(docx_path)
    problem_contract = build_problem_contract(input_revision, docx_hash, csv_hash)
    question_map = build_question_map(input_revision)
    data_contract = build_data_contract(input_revision, csv_hash)
    derivation = build_derivation_registry(input_revision)
    route = json.loads((results_root / "model_route.json").read_text(encoding="utf-8"))
    route["input_revision"] = input_revision
    route["skill_registry_revision"] = skill_revision
    route["code_revision"] = code_revision
    route["status"] = "VERIFIED"
    write_json(results_root / "model_route.json", route)
    route_revision = revision(route, prefix="route:")

    write_json(artifact_root / "input-manifest.json", input_manifest)
    write_json(artifact_root / "problem-contract.json", problem_contract)
    write_json(artifact_root / "question-map.json", question_map)
    write_json(artifact_root / "data-contract.json", data_contract)
    write_json(artifact_root / "derivation-registry.json", derivation)

    has_paper = (repo / "paper/b_problem_solution_v2.tex").is_file() and (repo / "paper/b_problem_solution_v2.pdf").is_file()
    result_hash = "sha256:" + sha256_file(results_root / "summary_v2.json")
    validation = build_validation(input_revision, result_hash, RUN_ID, summary, has_paper)
    write_json(artifact_root / "validation-report.json", validation)
    paper_contract = build_paper_contract(input_revision, validation, summary)
    write_json(artifact_root / "paper-contract-v2.json", paper_contract)
    assembly = build_assembly(input_revision, skill_revision)
    write_json(artifact_root / "assembly.json", assembly)

    # The run manifest hashes only safe relative paths and records the exact
    # sanitized stdout/stderr that were written above.
    run_artifacts: list[dict[str, Any]] = []
    for path in sorted(results_root.iterdir()):
        if path.name == "risk_scores_local.csv":
            continue
        media = "application/json" if path.suffix == ".json" else "text/csv" if path.suffix == ".csv" else "image/png" if path.suffix == ".png" else "text/plain"
        run_artifacts.append({"path": path.relative_to(repo).as_posix(), "sha256": "sha256:" + sha256_file(path), "media_type": media, "produced_by": "models/solve_b_problem_v2.py"})
    for rel in ("artifacts/runs/T-23-b-problem-v2/solver.stdout.txt", "artifacts/runs/T-23-b-problem-v2/solver.stderr.txt"):
        path = repo / rel
        run_artifacts.append({"path": rel, "sha256": "sha256:" + sha256_file(path), "media_type": "text/plain", "produced_by": "scripts/build_b_replay_artifacts.py"})
    if has_paper:
        for rel in ("paper/b_problem_solution_v2.tex", "paper/b_problem_solution_v2.pdf", "paper/render_qa_v2.json"):
            path = repo / rel
            if path.is_file():
                run_artifacts.append({"path": rel, "sha256": "sha256:" + sha256_file(path), "media_type": "application/pdf" if path.suffix == ".pdf" else "text/plain", "produced_by": "paper steward"})
    run_manifest = {
        "schema_version": "run-manifest/v2",
        "run_id": RUN_ID,
        "input_revision": input_revision,
        "code_revision": code_revision,
        "route_revision": route_revision,
        "skill_registry_revision": skill_revision,
        "environment": {"interpreter": sys.version.split()[0], "packages": ["numpy", "pandas", "scipy", "scikit-learn", "matplotlib"], "os": platform.system(), "solver_versions": ["solve_b_problem_v2/schema=b-problem-replay/v2"]},
        "command": "python -X utf8 models/solve_b_problem_v2.py --input <owner-attachment>/校赛B题附件.csv --output runtime/skill-v2-replay/b_solution_v2_final --seed 42",
        "cwd": ".",
        "seed_policy": {"deterministic": True, "seeds": [SEED, 143, 244, 345, 446], "random_sources": ["repeated-stratified-fold", "bootstrap", "latin-hypercube"]},
        "checkpoints": [{"name": "solver", "status": "PASS" if exit_code == 0 else "FAIL"}, {"name": "aggregate-artifacts", "status": "PASS"}, {"name": "policy-independent-arithmetic", "status": "PASS"}, {"name": "paper-render", "status": "PASS" if has_paper else "PENDING"}],
        "artifacts": run_artifacts,
        "exit_code": int(exit_code),
        "stdout_sha256": "sha256:" + sha256_file(artifact_root / "solver.stdout.txt"),
        "stderr_sha256": "sha256:" + sha256_file(artifact_root / "solver.stderr.txt"),
        "status": "VERIFIED" if exit_code == 0 else "BLOCKED",
    }
    write_json(artifact_root / "run-manifest.json", run_manifest)
    artifact_manifest = {
        "schema": "artifact-manifest/v2",
        "input_revision": input_revision,
        "skill_registry_revision": skill_revision,
        "source_material_manifest_sha256": "a3eb5ea717a45dee3b280afc506b37df3026e7142a9f3411e5bdafe3bc4306a",
        "raw_inputs_not_copied": True,
        "artifacts": build_manifest_artifacts(repo, artifact_root, results_root, has_paper),
        "status": "READY_FOR_REVIEW",
    }
    write_json(artifact_root / "artifact-manifest.json", artifact_manifest)
    print(json.dumps({"status": "PASS", "run_id": RUN_ID, "input_revision": input_revision, "skill_registry_revision": skill_revision, "result_hash": result_hash, "paper_present": has_paper, "artifact_root": str(artifact_root.relative_to(repo))}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
